import os
import uuid
from typing import List, Dict, Any, Optional
import pandas as pd
from docx import Document as DocxDocument
from datetime import datetime

from app.db.mongodb import documents_collection
from app.models.document import DocumentType
from app.schemas.document import DocumentCreate, DocumentContent

class DocumentProcessor:
    """Service for processing different types of documents and storing them in MongoDB"""
    
    @staticmethod
    async def process_document(document: DocumentCreate) -> str:
        """Process a document and store it in MongoDB"""
        document_id = str(uuid.uuid4())
        
        # Create MongoDB document
        mongo_doc = {
            "document_id": document_id,
            "org_id": str(document.org_id),
            "document_type": document.document_type.value,
            "content": document.content,
            "metadata": {
                "title": document.title,
                "description": document.description,
                "owner_email": document.owner_email,
                "tags": document.tags or [],
                "file_format": document.file_format
            },
            "version_history": [],
            "created_at": datetime.utcnow(),
            "updated_at": datetime.utcnow()
        }
        
        # Insert into MongoDB
        documents_collection.insert_one(mongo_doc)
        
        return document_id
    
    @staticmethod
    async def update_document(document_id: str, content: str, metadata: Optional[Dict[str, Any]] = None) -> bool:
        """Update a document in MongoDB"""
        # Get current document to save in version history
        current_doc = documents_collection.find_one({"document_id": document_id})
        if not current_doc:
            return False
        
        # Create version history entry
        version_entry = {
            "content": current_doc["content"],
            "metadata": current_doc["metadata"],
            "timestamp": datetime.utcnow()
        }
        
        # Update document
        update_data = {
            "content": content,
            "updated_at": datetime.utcnow(),
            "$push": {"version_history": version_entry}
        }
        
        if metadata:
            update_data["metadata"] = {**current_doc["metadata"], **metadata}
        
        result = documents_collection.update_one(
            {"document_id": document_id},
            {"$set": update_data}
        )
        
        return result.modified_count > 0
    
    @staticmethod
    async def get_document_content(document_id: str) -> Optional[DocumentContent]:
        """Get document content from MongoDB"""
        doc = documents_collection.find_one({"document_id": document_id})
        if not doc:
            return None
        
        return DocumentContent(
            document_id=doc["document_id"],
            content=doc["content"],
            metadata=doc["metadata"],
            version_history=doc["version_history"]
        )
    
    @staticmethod
    async def process_excel_file(file_path: str, document_type: DocumentType, org_id: str, owner_email: str) -> str:
        """Process an Excel file and store it in MongoDB"""
        # Read Excel file
        df = pd.read_excel(file_path)
        
        # Convert to JSON
        content = df.to_json(orient="records")
        
        # Create document
        document = DocumentCreate(
            title=os.path.basename(file_path),
            document_type=document_type,
            description=f"Excel file imported on {datetime.utcnow()}",
            org_id=org_id,
            owner_email=owner_email,
            content=content,
            file_format="excel"
        )
        
        return await DocumentProcessor.process_document(document)
    
    @staticmethod
    async def process_docx_file(file_path: str, document_type: DocumentType, org_id: str, owner_email: str) -> str:
        """Process a Word document and store it in MongoDB"""
        # Read Word document
        doc = DocxDocument(file_path)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        content = "\n".join(full_text)
        
        # Create document
        document = DocumentCreate(
            title=os.path.basename(file_path),
            document_type=document_type,
            description=f"Word document imported on {datetime.utcnow()}",
            org_id=org_id,
            owner_email=owner_email,
            content=content,
            file_format="docx"
        )
        
        return await DocumentProcessor.process_document(document)
    
    @staticmethod
    async def process_csv_file(file_path: str, document_type: DocumentType, org_id: str, owner_email: str) -> str:
        """Process a CSV file and store it in MongoDB"""
        # Read CSV file
        df = pd.read_csv(file_path)
        
        # Convert to JSON
        content = df.to_json(orient="records")
        
        # Create document
        document = DocumentCreate(
            title=os.path.basename(file_path),
            document_type=document_type,
            description=f"CSV file imported on {datetime.utcnow()}",
            org_id=org_id,
            owner_email=owner_email,
            content=content,
            file_format="csv"
        )
        
        return await DocumentProcessor.process_document(document)
